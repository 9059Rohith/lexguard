import io
import os
import re
import math
from typing import Optional, List, Dict, Any

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# ── Security: Strict upload directory allowlist ───────────────────────────────
ALLOWED_UPLOAD_BASE: Optional[str] = None  # Set at startup by main.py


def _safe_path(file_path: str, upload_dir: str) -> str:
    """
    Prevent path traversal. Resolves the file path and ensures it is
    strictly inside upload_dir. Raises ValueError for any traversal attempt.
    """
    resolved = os.path.realpath(os.path.abspath(file_path))
    base = os.path.realpath(os.path.abspath(upload_dir))
    if not resolved.startswith(base + os.sep) and resolved != base:
        raise ValueError(f"Path traversal attempt blocked: {file_path!r}")
    return resolved


class BoundingBox:
    """Normalized bounding box (x_min, y_min, x_max, y_max) in [0,1] space."""
    def __init__(self, x_min: float, y_min: float, x_max: float, y_max: float, page: int):
        self.x_min = round(max(0.0, min(1.0, x_min)), 4)
        self.y_min = round(max(0.0, min(1.0, y_min)), 4)
        self.x_max = round(max(0.0, min(1.0, x_max)), 4)
        self.y_max = round(max(0.0, min(1.0, y_max)), 4)
        self.page = page

    def to_dict(self) -> Dict[str, Any]:
        return {
            "x_min": self.x_min,
            "y_min": self.y_min,
            "x_max": self.x_max,
            "y_max": self.y_max,
            "page": self.page,
        }


class ParsedDocument:
    def __init__(self, full_text: str, page_count: int, chunks: list,
                 bounding_boxes: Optional[List[Dict]] = None,
                 counterparty: str = "", jurisdiction: str = ""):
        self.full_text = full_text
        self.page_count = page_count
        self.chunks = chunks  # list of {text, page, bbox?}
        self.bounding_boxes = bounding_boxes or []
        self.counterparty = counterparty
        self.jurisdiction = jurisdiction


# ── Counterparty / Jurisdiction Heuristics ────────────────────────────────────
_COUNTERPARTY_PATTERNS = [
    r'(?:between|entered into by and between)\s+([A-Z][A-Za-z\s,\.]+?)\s+(?:and|,)',
    r'(?:Client|Employer|Company|Lessor|Vendor|Licensor)[:\s]+([A-Z][A-Za-z\s,\.&]+?)(?:\n|,|\.|;)',
    r'(?:Counterparty|Party B|Second Party)[:\s]+([A-Z][A-Za-z\s,\.]+?)(?:\n|,)',
]
_JURISDICTION_PATTERNS = [
    r'(?:governed by|laws? of(?: the State of)?)\s+([A-Z][A-Za-z\s]+?)(?:\s+law|\s*[,\.\n])',
    r'(?:jurisdiction of|courts? of)\s+([A-Z][A-Za-z\s]+?)(?:\s*[,\.\n])',
    r'(?:State of|Commonwealth of)\s+([A-Z][A-Za-z]+)',
]


def _extract_counterparty(text: str) -> str:
    for pattern in _COUNTERPARTY_PATTERNS:
        m = re.search(pattern, text[:3000], re.IGNORECASE)
        if m:
            candidate = m.group(1).strip().rstrip(',.')
            if 5 <= len(candidate) <= 80:
                return candidate
    return ""


def _extract_jurisdiction(text: str) -> str:
    for pattern in _JURISDICTION_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            candidate = m.group(1).strip().rstrip(',.')
            if 3 <= len(candidate) <= 50:
                return candidate
    return ""


class DocumentParser:

    @staticmethod
    def parse(file_path: str, upload_dir: str = "./uploads") -> ParsedDocument:
        # Security: validate path
        try:
            safe = _safe_path(file_path, upload_dir)
        except ValueError:
            raise ValueError("File path validation failed — possible path traversal.")

        ext = os.path.splitext(safe)[1].lower()
        if ext == ".pdf":
            return DocumentParser._parse_pdf(safe)
        elif ext in (".docx", ".doc"):
            return DocumentParser._parse_docx(safe)
        elif ext in (".txt", ".md"):
            return DocumentParser._parse_txt(safe)
        elif ext == ".rtf":
            return DocumentParser._parse_rtf(safe)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _parse_pdf(file_path: str) -> ParsedDocument:
        if not HAS_FITZ:
            return DocumentParser._parse_txt_fallback(file_path)

        doc = fitz.open(file_path)
        full_text_parts = []
        chunks = []
        bboxes: List[Dict] = []
        page_count = doc.page_count

        for page_num in range(page_count):
            page = doc[page_num]
            pw, ph = page.rect.width, page.rect.height
            text = page.get_text("text")

            # OCR fallback for scanned pages
            if len(text.strip()) < 50 and HAS_OCR:
                try:
                    pix = page.get_pixmap(dpi=200)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    text = pytesseract.image_to_string(img)
                except Exception:
                    pass

            if text.strip():
                full_text_parts.append(text)

                # Extract word-level blocks for spatial coordinates
                blocks = page.get_text("blocks")  # (x0,y0,x1,y1,text,block_no,block_type)
                chunk_bboxes = []
                for block in blocks:
                    x0, y0, x1, y1, blk_text, *_ = block
                    if blk_text.strip():
                        bbox = BoundingBox(
                            x_min=x0 / pw if pw > 0 else 0,
                            y_min=y0 / ph if ph > 0 else 0,
                            x_max=x1 / pw if pw > 0 else 1,
                            y_max=y1 / ph if ph > 0 else 1,
                            page=page_num + 1,
                        )
                        chunk_bboxes.append(bbox.to_dict())
                        bboxes.append(bbox.to_dict())

                chunks.append({
                    "text": text,
                    "page": page_num + 1,
                    "bboxes": chunk_bboxes,
                })

        doc.close()
        full_text = "\n\n".join(full_text_parts)
        counterparty = _extract_counterparty(full_text)
        jurisdiction = _extract_jurisdiction(full_text)
        return ParsedDocument(
            full_text=full_text,
            page_count=page_count,
            chunks=chunks,
            bounding_boxes=bboxes,
            counterparty=counterparty,
            jurisdiction=jurisdiction,
        )

    @staticmethod
    def _parse_docx(file_path: str) -> ParsedDocument:
        if not HAS_DOCX:
            return ParsedDocument(full_text="Could not parse DOCX.", page_count=1, chunks=[])

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        word_count = len(full_text.split())
        page_count = max(1, word_count // 500)

        # Simulate normalized bounding boxes (evenly distributed)
        bboxes = []
        chunk_size = max(1, len(paragraphs) // max(1, page_count))
        for i, para in enumerate(paragraphs):
            page = (i // chunk_size) + 1
            y_pos = (i % chunk_size) / max(1, chunk_size)
            bbox = {
                "x_min": 0.05, "y_min": round(y_pos, 4),
                "x_max": 0.95, "y_max": round(min(1.0, y_pos + 0.06), 4),
                "page": page,
            }
            bboxes.append(bbox)

        counterparty = _extract_counterparty(full_text)
        jurisdiction = _extract_jurisdiction(full_text)
        return ParsedDocument(
            full_text=full_text,
            page_count=page_count,
            chunks=[{"text": full_text, "page": 1, "bboxes": bboxes}],
            bounding_boxes=bboxes,
            counterparty=counterparty,
            jurisdiction=jurisdiction,
        )

    @staticmethod
    def _parse_txt(file_path: str) -> ParsedDocument:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        word_count = len(text.split())
        page_count = max(1, word_count // 500)
        counterparty = _extract_counterparty(text)
        jurisdiction = _extract_jurisdiction(text)
        return ParsedDocument(
            full_text=text,
            page_count=page_count,
            chunks=[{"text": text, "page": 1}],
            counterparty=counterparty,
            jurisdiction=jurisdiction,
        )

    @staticmethod
    def _parse_rtf(file_path: str) -> ParsedDocument:
        """Parse RTF by stripping RTF control words and returning plain text."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        # Strip RTF control words: \word and \word123, and group braces
        text = re.sub(r'\\[a-z]+\-?\d*\s?', ' ', raw)
        text = re.sub(r'[{}]', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        word_count = len(text.split())
        page_count = max(1, word_count // 500)
        counterparty = _extract_counterparty(text)
        jurisdiction = _extract_jurisdiction(text)
        return ParsedDocument(
            full_text=text,
            page_count=page_count,
            chunks=[{"text": text, "page": 1}],
            counterparty=counterparty,
            jurisdiction=jurisdiction,
        )

    @staticmethod
    def _parse_txt_fallback(file_path: str) -> ParsedDocument:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception:
            text = "Could not parse document."
        return ParsedDocument(full_text=text, page_count=1, chunks=[{"text": text, "page": 1}])
